import docx
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
document = docx.Document()
sections = document.sections
for section in sections:
	# change orientation to landscape
	section.orientation = WD_ORIENT.LANDSCAPE

	new_width, new_height = section.page_height, section.page_width
	section.page_width = new_width
	section.page_height = new_height
document.add_picture('my_timetable/members/static/img/logo.png', width=Inches(6.25))
tab_header=["S.NO","year/sem","Date","TimeDuration","Register number","subcode","labvenue","Internal","External"]
table=document.add_table(rows=1,cols=9)
for i in range(9):
    table.rows[0].cells[i].text=tab_header[i]

frm=int(input("Enter the from reg no:"))
to=int(input("Enter the to reg no:"))
n=int(input("Enter the number of labs:"))
lbname=list(map(str,input("Enter the labname with a space:").strip().split()))[:n]
r=int(input("Enter the capacity of each lab:"))
sem=input("Enter the year and semester(year/sem):")
fd=int(input("Enter from date:"))
fm=int(input("Enter from month:"))
fy=int(input("Enter from year:"))
t1="8.30am-11.30am"
t2="12.00no-3.00pm"
l=int(input("Enter the number of subjects:"))
subcode=list(map(str,input("Enter the subject code separated by a space:").strip().split()))[:l]
internal=list(map(str,input("Enter the names of the internal examiner with a space:").strip().split()))[:n]
external=list(map(str,input("Enter the names of the External examiner with a space:").strip().split()))[:n]
time=1
i=0
sno=1

c=to-frm
while(1):
    if(c%r==0 or c%r>=15) :
        c=r
        break
    elif(c%r<15):
        r=r+1

if (sem == "I/II" or sem == "I/I"):
    num = to - frm
    ro = 2 * (math.ceil(num / c))
    ro1 = ro / 2
    for pq in range(len(fd)):
        i = 0
        frmc = frm
        toc = to
        for ij in range(ro):
            cells = table.add_rows().cells
            for j in range(10):
                if (j == 0):
                    cells[0].text = str(sno)
                    sno = sno + 1
                elif (j == 1):
                    cells[1].text = sem
                elif (j == 2):
                    k = fd[ii]
                    cells[2].text = k
                elif (j == 3):
                    cells[3].text = t3[kk]
                    kk = (kk + 1) % 4
                elif (j == 4):
                    if ((frmc <= toc) and (frmc + c + 1 <= toc)):
                        p = str(frmc)
                        u = str(frmc + c - 1)
                        cells[4].text = p
                        cells[5].text = u
                    else:
                        p = str(frmc)
                        u = str(toc)
                        cells[4].text = p
                        cells[5].text = u
                elif (j == 6):
                    if (ij <= ro1):
                        cells[6].text = subcode[0]
                    else:
                        cells[6].text = subcode[1]

                elif (j == 7):
                    cells[7].text = lbname[i]
                elif (j == 8):
                    cells[8].text = internal[i]
                elif (j == 9):
                    cells[9].text = external[i]
            frmc = frmc + c
            print("Hello")
        ii = ii + 1
        table.add_row().cells
    table.style = 'Table Grid'
    document.save("text1.docx")



