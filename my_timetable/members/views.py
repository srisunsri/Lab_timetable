import docx
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from django.http import HttpResponse
from django.shortcuts import render
from django.views.decorators.csrf import ensure_csrf_cookie
from docxtpl import DocxTemplate
@ensure_csrf_cookie
def members(request):
  return render(request,'index.html')

def product(request):
  frm = int(request.POST["fromreg"])
  to = int(request.POST["toreg"])
  n = int(request.POST["nooflab"])
  lbname = list(map(str, request.POST["labname"].strip().split()))[:n]
  r = int(request.POST["capacity"])
  sem = request.POST["year"]
  fd = request.POST["date"].strip().split()
  t1 = "8.30am-11.30am"
  t2 = "12.00no-3.00pm"
  t3 = ["8.30am-10.00am","10.00am-11.30am","12.00no-1.30pm","1.30pm-3.00pm"]
  l = int(request.POST["subjects"])
  subcode = list(map(str, request.POST["subcode"].strip().split()))
  internal = list(map(str, request.POST["internal"].strip().split()))[:n]
  external = list(map(str, request.POST["external"].strip().split()))[:n]
  sub = list(map(str, request.POST["sub"].strip().split()))
  document = docx.Document()
  sections = document.sections
  for section in sections:
    # change orientation to landscape
    section.orientation = WD_ORIENT.LANDSCAPE

    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
  document.add_picture('members/static/img/logo.png', width=Inches(6.25))
  tab_header = ["S.NO", "year/sem", "Date", "TimeDuration", "From Reg.No", "To Reg.No", "subcode", "labvenue",
                "Internal",
                "External"]
  table = document.add_table(rows=1, cols=10)
  for i in range(10):
    table.rows[0].cells[i].text = tab_header[i]

  c = to - frm
  while (1):
    if (c % r == 0 or c % r >= 15):
      c = r
      break
    elif (c % r < 15):
      r = r + 1
  time = 1
  i = 0
  sno = 1
  ii = 0
  for z in subcode:
    i = 0
    frmc = frm
    toc = to
    while (frmc <= toc):
      cells = table.add_row().cells
      for j in range(10):
        if (i >= n):
          i = 0
        if (j == 0):
          cells[0].text = str(sno)
          sno = sno + 1
        elif (j == 1):
          cells[1].text = sem
        elif (j == 2):
          k = fd[ii]
          cells[2].text = k
        elif (j == 3):
          if (time % 2 == 1):
            cells[3].text = t1
            time = 0
          else:
            cells[3].text = t2
            time = 1
        elif (j == 4):
          if ((frmc <= toc) and (frmc + c + 1 <= toc)):
            p = str(frmc)
            u = str(frmc + c - 1)
            cells[4].text = p
            cells[5].text = u

          else:
            p = str(frmc)
            u = str(toc)
            cells[4].text = p
            cells[5].text = u
        elif (j == 6):
          cells[6].text = z
        elif (j == 7):
          cells[7].text = lbname[i]
        elif (j == 8):
          cells[8].text = internal[i]
        elif (j == 9):
          cells[9].text = external[i]
        print("H", end=" ")
      if (time == 1):
        i = i + 1
      frmc = frmc + c
      print("Hello")
    ii = ii + 1
    table.add_row().cells
  go = 0
  time = 1
  tim = 1
  for z in sub:
    i = 0
    frmc = frm
    toc = to
    while (frmc <= toc):
      cells = table.add_row().cells
      for j in range(10):
        if (i >= n):
          i = 0
        if (j == 0):
          cells[0].text = str(sno)
          sno = sno + 1
        elif (j == 1):
          cells[1].text = sem
        elif (j == 2):
          k = fd[ii]
          cells[2].text = k
        elif (j == 3):
          cells[3].text = t3[go]
          go = (go + 1) % 4
        elif (j == 4):
          if ((frmc <= toc) and (frmc + c + 1 <= toc)):
            p = str(frmc)
            u = str(frmc + c - 1)
            cells[4].text = p
            cells[5].text = u
            if(time==0):
              time=1
            else:
              time=0

          else:
            p = str(frmc)
            u = str(toc)
            cells[4].text = p
            cells[5].text = u
            if (time == 0):
              time = 1
            else:
              time = 0
        elif (j == 6):
          cells[6].text = z
        elif (j == 7):
          cells[7].text = lbname[i]
        elif (j == 8):
          cells[8].text = internal[i]
        elif (j == 9):
          cells[9].text = external[i]
      print("sri")
      if (time == 1):
        frmc = frmc + c
      if (tim == 1):
        i = i + 1
        tim = 0
      elif (tim ==0):
        i = i - 1
        tim=
    print("sun")
    ii = ii + 1
    table.add_row().cells
  table.style = 'Table Grid'
  document.save("text1.docx")

  document = docx.Document()
  sections = document.sections
  for section in sections:
    # change orientation to landscape
    section.orientation = WD_ORIENT.LANDSCAPE

    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
  document.add_picture('members/static/img/logo.png', width=Inches(6.25))
  tab_header = ["S.NO", "year/sem", "Date", "TimeDuration", "From Reg.No", "To Reg.No", "subcode", "labvenue"]
  table = document.add_table(rows=1, cols=8)
  for i in range(8):
    table.rows[0].cells[i].text = tab_header[i]
  time = 1
  i = 0
  sno = 1

  c = to - frm
  while (1):
    if (c % r == 0 or c % r >= 15):
      c = r
      break
    elif (c % r < 15):
      r = r + 1
  ii = 0
  for z in subcode:
    i = 0
    frmc = frm
    toc = to
    while (frmc <= toc):
      cells = table.add_row().cells
      for j in range(8):
        if (i >= n):
          i = 0
        if (j == 0):
          cells[0].text = str(sno)
          sno = sno + 1
        elif (j == 1):
          cells[1].text = sem
        elif (j == 2):
          k = fd[ii]
          cells[2].text = k
        elif (j == 3):
          if (time % 2 == 1):
            cells[3].text = t1
            time = 0
          else:
            cells[3].text = t2
            time = 1
        elif (j == 4):
          if ((frmc <= toc) and (frmc + c + 1 <= toc)):
            p = str(frmc)
            u = str(frmc + c - 1)
            cells[4].text = p
            cells[5].text = u

          else:
            p = str(frmc)
            u = str(toc)
            cells[4].text = p
            cells[5].text = u
        elif (j == 6):
          cells[6].text = z
        elif (j == 7):
          cells[7].text = lbname[i]

      if (time == 1):
        i = i + 1
      frmc = frmc + c
      print("Hello")
    ii = ii + 1
    table.add_row().cells
  table.style = 'Table Grid'

  document.save("text.docx")
  return render(request, 'result.html')




def render_docx(request):
  doc = DocxTemplate("text1.docx")
  # you have to place your_docx_template.docx in the root of your project (same level as manage.py).

  context = {
    # ...
  }

  response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
  response["Content-Disposition"] = 'filename="your_doc_name.docx"'

  doc.render(context)
  doc.save(response)

  return response

def renderdocx(request):
  doc = DocxTemplate("text.docx")
  # you have to place your_docx_template.docx in the root of your project (same level as manage.py).

  context = {
    # ...
  }

  response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
  response["Content-Disposition"] = 'filename="forstudents.docx"'

  doc.render(context)
  doc.save(response)

  return response